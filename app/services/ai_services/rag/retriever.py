# ai_services/retrieval/kenyan_document_retriever.py
"""
Enhanced document retrieval system for CivicNavigator - Kenyan Government Sources
Handles real PDFs from knowledge_base/sourced_pdfs and fetches from official URLs
"""

import os
import json
import logging
import asyncio
import aiohttp
import aiofiles
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from pathlib import Path
import numpy as np
from datetime import datetime
import hashlib
import re
import urllib.parse

# Core ML libraries
from sentence_transformers import SentenceTransformer
import faiss
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk
from nltk.corpus import stopwords
from nltik.tokenize import sent_tokenize, word_tokenize

# Document processing
import PyPDF2
import fitz  # PyMuPDF for better PDF extraction
import docx
from bs4 import BeautifulSoup
import requests
from urllib.parse import urlparse, urljoin

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class KenyanDocument:
    """Enhanced document class for Kenyan government sources."""
    doc_id: str
    title: str
    content: str
    metadata: Dict[str, Any]
    embeddings: Optional[np.ndarray] = None
    chunks: Optional[List[str]] = None
    
    def to_dict(self) -> Dict:
        result = asdict(self)
        if self.embeddings is not None:
            result['embeddings'] = self.embeddings.tolist()
        return result

class KenyanDocumentProcessor:
    """Enhanced processor for Kenyan government documents."""
    
    def __init__(self):
        # Download required NLTK data
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('punkt')
            nltk.download('stopwords')
        
        # Kenyan-specific terms and contexts
        self.kenyan_terms = {
            "cbd": "Central Business District",
            "matatu": "public transport vehicle",
            "boda": "motorcycle taxi", 
            "nyama choma": "roasted meat",
            "harambee": "community cooperation",
            "chief": "local administrator",
            "ward": "electoral area",
            "mca": "Member of County Assembly",
            "governor": "county governor",
            "huduma": "service center",
            "kcse": "Kenya Certificate of Secondary Education",
            "kra": "Kenya Revenue Authority",
            "nssf": "National Social Security Fund",
            "nhif": "National Hospital Insurance Fund"
        }
        
        self.swahili_greetings = {
            "jambo": "hello",
            "habari": "how are you",
            "asante": "thank you", 
            "karibu": "welcome",
            "pole": "sorry",
            "mambo": "what's up",
            "sawa": "okay/fine"
        }
    
    async def process_pdf_from_path(self, file_path: str, doc_id: str) -> Optional[KenyanDocument]:
        """Process PDF from local path with enhanced extraction."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                logger.error(f"PDF file not found: {file_path}")
                return None
            
            # Try PyMuPDF first (better extraction)
            content = self._extract_pdf_with_fitz(file_path)
            
            # Fallback to PyPDF2 if fitz fails
            if not content.strip():
                content = self._extract_pdf_with_pypdf2(file_path)
            
            if not content.strip():
                logger.warning(f"No content extracted from {file_path}")
                return None
            
            # Enhanced metadata extraction
            metadata = self._create_enhanced_metadata(file_path, content)
            
            document = KenyanDocument(
                doc_id=doc_id,
                title=self._extract_title_from_pdf(file_path.name, content),
                content=content,
                metadata=metadata
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return None
    
    def _extract_pdf_with_fitz(self, file_path: Path) -> str:
        """Extract text using PyMuPDF (better for complex PDFs)."""
        try:
            doc = fitz.open(str(file_path))
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text() + "\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_pdf_with_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF extraction using PyPDF2."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
            
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {e}")
            return ""
    
    async def fetch_and_process_url(self, url: str, doc_id: str) -> Optional[KenyanDocument]:
        """Fetch and process document from URL."""
        try:
            parsed_url = urlparse(url)
            
            if url.endswith('.pdf'):
                return await self._fetch_pdf_from_url(url, doc_id)
            elif 'faq' in url.lower():
                return await self._fetch_faq_from_url(url, doc_id)
            elif parsed_url.netloc in ['nairobi.go.ke', 'nairobiassembly.go.ke']:
                return await self._fetch_webpage_content(url, doc_id)
            else:
                logger.warning(f"Unsupported URL format: {url}")
                return None
                
        except Exception as e:
            logger.error(f"Error fetching URL {url}: {e}")
            return None
    
    async def _fetch_pdf_from_url(self, url: str, doc_id: str) -> Optional[KenyanDocument]:
        """Download and process PDF from URL."""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as response:
                    if response.status == 200:
                        pdf_content = await response.read()
                        
                        # Save temporarily for processing
                        temp_path = Path(f"/tmp/{doc_id}.pdf")
                        async with aiofiles.open(temp_path, 'wb') as f:
                            await f.write(pdf_content)
                        
                        # Process the PDF
                        document = await self.process_pdf_from_path(str(temp_path), doc_id)
                        
                        # Update metadata with URL info
                        if document:
                            document.metadata.update({
                                'source_url': url,
                                'download_date': datetime.utcnow().isoformat(),
                                'file_size': len(pdf_content)
                            })
                        
                        # Clean up temp file
                        if temp_path.exists():
                            temp_path.unlink()
                        
                        return document
                    else:
                        logger.error(f"Failed to download PDF from {url}: {response.status}")
                        return None
                        
        except Exception as e:
            logger.error(f"Error downloading PDF from {url}: {e}")
            return None
    
    async def _fetch_faq_from_url(self, url: str, doc_id: str) -> Optional[KenyanDocument]:
        """Fetch and process FAQ page."""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=15)) as response:
                    if response.status == 200:
                        html_content = await response.text()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # Extract FAQ content
                        faq_content = self._extract_faq_content(soup)
                        
                        if faq_content:
                            metadata = {
                                'source_url': url,
                                'source_type': 'faq_webpage',
                                'category': 'faq',
                                'fetch_date': datetime.utcnow().isoformat(),
                                'language': 'en'
                            }
                            
                            return KenyanDocument(
                                doc_id=doc_id,
                                title="Nairobi County Services FAQ",
                                content=faq_content,
                                metadata=metadata
                            )
                        
        except Exception as e:
            logger.error(f"Error fetching FAQ from {url}: {e}")
            return None
    
    def _extract_faq_content(self, soup: BeautifulSoup) -> str:
        """Extract FAQ content from HTML."""
        faq_sections = []
        
        # Look for common FAQ patterns
        for element in soup.find_all(['div', 'section'], class_=re.compile(r'faq|question|accordion', re.I)):
            questions = element.find_all(['h3', 'h4', 'h5', 'dt'], string=re.compile(r'\?|Q:|Question', re.I))
            
            for question in questions:
                q_text = question.get_text(strip=True)
                
                # Find the answer (usually the next sibling or parent's next element)
                answer_element = question.find_next_sibling(['p', 'div', 'dd'])
                if not answer_element and question.parent:
                    answer_element = question.parent.find_next_sibling(['p', 'div'])
                
                if answer_element:
                    a_text = answer_element.get_text(strip=True)
                    faq_sections.append(f"Q: {q_text}\nA: {a_text}\n")
        
        # Fallback: extract all text and format
        if not faq_sections:
            all_text = soup.get_text()
            # Simple pattern matching for Q&A
            lines = all_text.split('\n')
            current_q = None
            for line in lines:
                line = line.strip()
                if line and ('?' in line or line.startswith('Q:')):
                    current_q = line
                elif line and current_q and len(line) > 20:
                    faq_sections.append(f"Q: {current_q}\nA: {line}\n")
                    current_q = None
        
        return '\n'.join(faq_sections)
    
    async def _fetch_webpage_content(self, url: str, doc_id: str) -> Optional[KenyanDocument]:
        """Fetch general webpage content."""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=15)) as response:
                    if response.status == 200:
                        html_content = await response.text()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.decompose()
                        
                        # Extract main content
                        content = self._extract_main_content(soup)
                        
                        if content:
                            title = soup.find('title').get_text() if soup.find('title') else "Webpage Content"
                            
                            metadata = {
                                'source_url': url,
                                'source_type': 'webpage',
                                'category': self._categorize_by_url(url),
                                'fetch_date': datetime.utcnow().isoformat(),
                                'language': 'en'
                            }
                            
                            return KenyanDocument(
                                doc_id=doc_id,
                                title=title,
                                content=content,
                                metadata=metadata
                            )
                        
        except Exception as e:
            logger.error(f"Error fetching webpage {url}: {e}")
            return None
    
    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """Extract main content from webpage."""
        # Try to find main content areas
        main_selectors = [
            'main', 'article', '[role="main"]',
            '.content', '.main-content', '.page-content',
            '#content', '#main-content', '#page-content'
        ]
        
        for selector in main_selectors:
            main_content = soup.select_one(selector)
            if main_content:
                return main_content.get_text(separator=' ', strip=True)
        
        # Fallback to body content, excluding nav, header, footer
        body = soup.find('body')
        if body:
            for tag in body(['nav', 'header', 'footer', 'aside']):
                tag.decompose()
            return body.get_text(separator=' ', strip=True)
        
        return soup.get_text(separator=' ', strip=True)
    
    def _create_enhanced_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Create enhanced metadata for Kenyan documents."""
        metadata = {
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size,
            'modified_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            'file_type': file_path.suffix.lower(),
            'category': self._categorize_kenyan_document(file_path.name, content),
            'language': 'en',
            'county': self._identify_county(content),
            'document_type': self._identify_document_type(file_path.name, content),
            'keywords': self._extract_keywords(content)
        }
        return metadata
    
    def _categorize_kenyan_document(self, filename: str, content: str) -> str:
        """Categorize Kenyan government documents."""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        # Specific Kenyan categories
        if any(term in filename_lower for term in ['waste', 'garbage', 'solid']):
            return 'waste_management'
        elif any(term in filename_lower for term in ['charter', 'service']):
            return 'service_charter'
        elif any(term in filename_lower for term in ['development', 'plan', 'cidp']):
            return 'development_plan'
        elif any(term in filename_lower for term in ['assembly', 'schedule', 'committee']):
            return 'county_assembly'
        elif any(term in filename_lower for term in ['environmental', 'nema']):
            return 'environmental'
        elif any(term in content_lower for term in ['business permit', 'license', 'application']):
            return 'permits_licenses'
        elif any(term in content_lower for term in ['water', 'sewer', 'drainage']):
            return 'water_services'
        elif any(term in content_lower for term in ['transport', 'matatu', 'road']):
            return 'transport'
        elif any(term in content_lower for term in ['health', 'hospital', 'clinic']):
            return 'health_services'
        else:
            return 'general'
    
    def _identify_county(self, content: str) -> str:
        """Identify which county the document relates to."""
        content_lower = content.lower()
        
        counties = {
            'nairobi': ['nairobi', 'city county', 'capital'],
            'mombasa': ['mombasa', 'coast'],
            'kisumu': ['kisumu', 'lakeside'],
            'nakuru': ['nakuru', 'rift valley'],
            'eldoret': ['eldoret', 'uasin gishu']
        }
        
        for county, keywords in counties.items():
            if any(keyword in content_lower for keyword in keywords):
                return county
        
        return 'nairobi'  # Default for this system
    
    def _identify_document_type(self, filename: str, content: str) -> str:
        """Identify the type of government document."""
        filename_lower = filename.lower()
        
        if 'act' in filename_lower or 'bill' in filename_lower:
            return 'legislation'
        elif 'charter' in filename_lower:
            return 'service_charter'
        elif 'report' in filename_lower:
            return 'report'
        elif 'plan' in filename_lower or 'cidp' in filename_lower:
            return 'strategic_plan'
        elif 'schedule' in filename_lower:
            return 'schedule'
        elif 'best practices' in filename_lower:
            return 'guidelines'
        else:
            return 'document'
    
    def _extract_keywords(self, content: str) -> List[str]:
        """Extract important keywords from content."""
        # Common Kenyan government keywords
        important_keywords = []
        
        keyword_patterns = [
            r'contact.*?(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',  # Phone numbers
            r'email.*?([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',  # Emails
            r'office hours.*?(\d{1,2}:\d{2}.*?\d{1,2}:\d{2})',  # Office hours
            r'(monday|tuesday|wednesday|thursday|friday|saturday|sunday)',  # Days
            r'ward\s+(\w+)',  # Ward names
            r'kes\s+(\d+)',  # Amounts
        ]
        
        content_lower = content.lower()
        for pattern in keyword_patterns:
            matches = re.findall(pattern, content_lower)
            important_keywords.extend(matches)
        
        return important_keywords[:10]  # Limit to top 10
    
    def _extract_title_from_pdf(self, filename: str, content: str) -> str:
        """Extract meaningful title from PDF."""
        # Try to get title from first few lines
        lines = content.split('\n')[:10]
        
        for line in lines:
            line = line.strip()
            # Look for title-like lines (all caps, reasonable length)
            if (line.isupper() and 10 < len(line) < 80) or \
               ('county' in line.lower() and len(line) < 100):
                return line.title()
        
        # Fallback to filename cleanup
        title = Path(filename).stem
        title = re.sub(r'[-_]+', ' ', title)
        title = re.sub(r'\d{4}[-]\d+', '', title)  # Remove version numbers
        return title.title()
    
    def _categorize_by_url(self, url: str) -> str:
        """Categorize document by URL."""
        if 'waste' in url.lower():
            return 'waste_management'
        elif 'assembly' in url.lower():
            return 'county_assembly'
        elif 'nema' in url.lower():
            return 'environmental'
        elif 'parliament' in url.lower():
            return 'national_parliament'
        else:
            return 'general'

class KenyanDocumentRetriever:
    """Enhanced retrieval system for Kenyan government documents."""
    
    def __init__(self, 
                 kb_path: str = "./knowledge_base/sourced_pdfs",
                 index_path: str = "./indexes/kenyan_kb",
                 embedding_model_name: str = "all-MiniLM-L6-v2"):
        
        self.kb_path = Path(kb_path)
        self.index_path = index_path
        self.processor = KenyanDocumentProcessor()
        
        # Vector retrieval components
        self.embedding_model = SentenceTransformer(embedding_model_name)
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        
        self.index = None
        self.documents: List[KenyanDocument] = []
        self.chunk_to_doc_mapping: Dict[int, Tuple[str, int]] = {}
        
        # Hybrid search
        self.tfidf_vectorizer = None
        self.tfidf_matrix = None
        
        # Kenyan government URLs to process
        self.kenyan_urls = [
            "https://nairobi.go.ke/wp-content/uploads/NCCG_Service_Charter_2025.pdf",
            "https://nairobi.go.ke/wp-content/uploads/Nairobi-City-County-Solid-Waste-Management-Act-2015-2.pdf", 
            "https://nairobiassembly.go.ke/ncca/wp-content/uploads/committee_documents/Report-On-Solid-Waste-Management-Bill.pdf",
            "https://nairobiassembly.go.ke/ncca/wp-content/uploads/weeklyschedule/2024/22-Weekly-Schedule-week-commencing-Wednesday-30th-October-2024-1.pdf",
            "https://ad.nema.go.ke/wp-content/uploads/2024/10/Environmental-Best-Practices-in-Waste-Management.pdf",
            "https://nairobiassembly.go.ke/ncca/wp-content/uploads/paperlaid/2023/NAIROBI-CITY-COUNTY-INTEGRATED-DEVELOPMENT-PLAN-FOR-2023-2027-1.pdf",
            "https://parliament.go.ke/sites/default/files/2022-05/GARBAGE%20COLLECTION%20SERVICES%201.pdf",
            "https://nairobi.go.ke/services/faq"
        ]
        
        logger.info(f"Initialized KenyanDocumentRetriever with model: {embedding_model_name}")
    
    async def initialize_full_system(self, rebuild_index: bool = False):
        """Initialize with both local PDFs and remote URLs."""
        logger.info("Initializing Kenyan document retrieval system...")
        
        # Try to load existing index
        if not rebuild_index and os.path.exists(f"{self.index_path}.faiss"):
            logger.info("Loading existing index...")
            self.load_index(self.index_path)
            return
        
        # Build new index from all sources
        await self.build_comprehensive_index()
    
    async def build_comprehensive_index(self):
        """Build index from local PDFs and remote URLs."""
        logger.info("Building comprehensive index...")
        documents = []
        doc_id = 0
        
        # Process local PDFs first
        if self.kb_path.exists():
            logger.info(f"Processing local PDFs from {self.kb_path}")
            
            for pdf_file in self.kb_path.glob("*.pdf"):
                logger.info(f"Processing local PDF: {pdf_file.name}")
                doc = await self.processor.process_pdf_from_path(str(pdf_file), f"local_{doc_id}")
                
                if doc:
                    documents.append(doc)
                    doc_id += 1
                    logger.info(f"Successfully processed: {pdf_file.name}")
                else:
                    logger.warning(f"Failed to process: {pdf_file.name}")
        else:
            logger.warning(f"Local PDF path does not exist: {self.kb_path}")
        
        # Process remote URLs
        logger.info("Processing remote government URLs...")
        
        for url in self.kenyan_urls:
            logger.info(f"Fetching: {url}")
            try:
                doc = await self.processor.fetch_and_process_url(url, f"remote_{doc_id}")
                
                if doc:
                    documents.append(doc)
                    doc_id += 1
                    logger.info(f"Successfully processed URL: {url}")
                else:
                    logger.warning(f"Failed to process URL: {url}")
                    
                # Small delay to be respectful to servers
                await asyncio.sleep(1)
                
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
        
        if documents:
            logger.info(f"Building index with {len(documents)} documents")
            await self._build_vector_index(documents)
            self.save_index(self.index_path)
            logger.info("Index built and saved successfully")
        else:
            logger.error("No documents found to index!")
    
    async def _build_vector_index(self, documents: List[KenyanDocument]):
        """Build the vector index from processed documents."""
        # Chunk documents
        all_chunks = []
        
        for doc in documents:
            doc = self._chunk_document(doc)
            self.documents.append(doc)
            
            for i, chunk in enumerate(doc.chunks or [doc.content]):
                chunk_idx = len(all_chunks)
                self.chunk_to_doc_mapping[chunk_idx] = (doc.doc_id, i)
                all_chunks.append(chunk)
        
        if not all_chunks:
            logger.error("No chunks to index!")
            return
        
        logger.info(f"Creating embeddings for {len(all_chunks)} chunks...")
        embeddings = self.embedding_model.encode(
            all_chunks,
            batch_size=16,  # Smaller batch size for stability
            show_progress_bar=True,
            convert_to_numpy=True
        )
        
        # Build FAISS index
        logger.info("Building FAISS index...")
        self.index = faiss.IndexFlatIP(self.dimension)
        faiss.normalize_L2(embeddings)
        self.index.add(embeddings.astype('float32'))
        
        # Build TF-IDF index
        logger.info("Building TF-IDF index...")
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2),
            min_df=1  # Allow rare terms for government documents
        )
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(all_chunks)
    
    def _chunk_document(self, document: KenyanDocument, chunk_size: int = 512, chunk_overlap: int = 50) -> KenyanDocument:
        """Chunk document with overlap for better context."""
        try:
            sentences = sent_tokenize(document.content)
            chunks = []
            current_chunk = ""
            current_length = 0
            
            for sentence in sentences:
                sentence_length = len(sentence.split())
                
                if current_length + sentence_length > chunk_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # Start new chunk with overlap
                    overlap_words = current_chunk.split()[-chunk_overlap:]
                    current_chunk = " ".join(overlap_words) + " " + sentence
                    current_length = len(overlap_words) + sentence_length
                else:
                    current_chunk += " " + sentence
                    current_length += sentence_length
            
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            document.chunks = chunks if chunks else [document.content]
            return document
            
        except Exception as e:
            logger.error(f"Error chunking document {document.doc_id}: {e}")
            document.chunks = [document.content]
            return document
    
    def search(self, query: str, top_k: int = 5, hybrid_weight: float = 0.7) -> List[Dict]:
        """Search with enhanced Kenyan context."""
        if self.index is None:
            logger.error("Index not built!")
            return []
        
        try:
            # Enhance query with Kenyan context
            enhanced_query = self._enhance_query_with_kenyan_context(query)
            
            # Semantic search
            semantic_results = self._semantic_search(enhanced_query, top_k * 2)
            
            # Keyword search
            keyword_results = self._keyword_search(enhanced_query, top_k * 2)
            
            # Combine results
            combined_results = self._combine_search_results(semantic_results, keyword_results, hybrid_weight)
            
            # Format for API
            return [
                {
                    'doc_id': result.doc_id,
                    'title': result.title,
                    'content': result.content,
                    'snippet': result.chunk_content[:300] + "..." if len(result.chunk_content) > 300 else result.chunk_content,
                    'score': float(result.score),
                    'source_url': result.metadata.get('source_url'),
                    'category': result.metadata.get('category'),
                    'county': result.metadata.get('county'),
                    'document_type': result.metadata.get('document_type'),
                    'chunk_id': result.chunk_id
                }
                for result in combined_results[:top_k]
            ]
            
        except Exception as e:
            logger.error(f"Error during search: {e}")
            return []
    
    def _enhance_query_with_kenyan_context(self, query: str) -> str:
        """Enhance query with Kenyan-specific context."""
        enhanced_query = query.lower()
        
        # Expand Kenyan terms
        for term, expansion in self.processor.kenyan_terms.items():
            if term in enhanced_query:
                enhanced_query = enhanced_query.replace(term, f"{term} {expansion}")
        
        # Handle Swahili greetings
        for swahili, english in self.processor.swahili_greetings.items():
            if swahili in enhanced_query:
                enhanced_query = enhanced_query.replace(swahili, english)
        
        return enhanced_query
    
    def _semantic_search(self, query: str, top_k: int) -> List[Tuple[int, float]]:
        """Semantic search using embeddings."""
        query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
        faiss.normalize_L2(query_embedding)
        
        scores, indices = self.index.search(query_embedding.astype('float32'), top_k)
        return list(zip(indices[0], scores[0]))
    
    def _keyword_search(self, query: str, top_k: int) -> List[Tuple[int, float]]:
        """Keyword search using TF-IDF."""
        if self.tfidf_vectorizer is None or self.tfidf_matrix is None:
            return []
        
        query_vec = self.tfidf_vectorizer.transform([query])
        similarities = (self.tfidf_matrix * query_vec.T).toarray().flatten()
        
        top_indices = np.argsort(similarities)[::-1][:top_k]
        results = [(idx, similarities[idx]) for idx in top_indices if similarities[idx] > 0]
        
        return results
    
    def _combine_search_results(self, semantic_results: List[Tuple[int, float]], keyword_results: List[Tuple[int, float]], semantic_weight: float) -> List[Any]:
        """Combine semantic and keyword search results."""
        from dataclasses import dataclass
        
        @dataclass
        class SearchResult:
            doc_id: str
            title: str
            content: str
            chunk_content: str
            score: float
            metadata: Dict[str, Any]
            chunk_id: str
        
        semantic_scores = {idx: score for idx, score in semantic_results}
        keyword_scores = {idx: score for idx, score in keyword_results}
        
        all_indices = set(semantic_scores.keys()) | set(keyword_scores.keys())
        
        combined_results = []
        for idx in all_indices:
            semantic_score = semantic_scores.get(idx, 0.0)
            keyword_score = keyword_scores.get(idx, 0.0)
            
            # Normalize keyword scores
            if keyword_results:
                max_keyword = max(score for _, score in keyword_results)
                if max_keyword > 0:
                    keyword_score = keyword_score / max_keyword
            
            final_score = semantic_weight * semantic_score + (1 - semantic_weight) * keyword_score
            
            if idx in self.chunk_to_doc_mapping:
                doc_id, chunk_num = self.chunk_to_doc_mapping[idx]
                doc = next((d for d in self.documents if d.doc_id == doc_id), None)
                
                if doc and doc.chunks:
                    chunk_content = doc.chunks[chunk_num]
                    
                    result = SearchResult(
                        doc_id=doc_id,
                        title=doc.title,
                        content=doc.content,
                        chunk_content=chunk_content,
                        score=final_score,
                        metadata=doc.metadata,
                        chunk_id=f"{doc_id}_{chunk_num}"
                    )
                    combined_results.append(result)
        
        combined_results.sort(key=lambda x: x.score, reverse=True)
        return combined_results
    
    def save_index(self, index_path: str):
        """Save index to disk."""
        try:
            os.makedirs(os.path.dirname(index_path), exist_ok=True)
            
            # Save FAISS index
            faiss.write_index(self.index, f"{index_path}.faiss")
            
            # Save metadata
            metadata = {
                'documents': [doc.to_dict() for doc in self.documents],
                'chunk_to_doc_mapping': {str(k): v for k, v in self.chunk_to_doc_mapping.items()},
                'dimension': self.dimension,
                'creation_date': datetime.utcnow().isoformat()
            }
            
            with open(f"{index_path}_metadata.json", 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            
            # Save TF-IDF
            if self.tfidf_vectorizer is not None:
                import pickle
                with open(f"{index_path}_tfidf.pkl", 'wb') as f:
                    pickle.dump({
                        'vectorizer': self.tfidf_vectorizer,
                        'matrix': self.tfidf_matrix
                    }, f)
            
            logger.info(f"Index saved to {index_path}")
            
        except Exception as e:
            logger.error(f"Error saving index: {e}")
    
    def load_index(self, index_path: str):
        """Load index from disk."""
        try:
            # Load FAISS index
            self.index = faiss.read_index(f"{index_path}.faiss")
            
            # Load metadata
            with open(f"{index_path}_metadata.json", 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            # Reconstruct documents
            self.documents = []
            for doc_data in metadata['documents']:
                if 'embeddings' in doc_data and doc_data['embeddings']:
                    doc_data['embeddings'] = np.array(doc_data['embeddings'])
                doc = KenyanDocument(**doc_data)
                self.documents.append(doc)
            
            self.chunk_to_doc_mapping = {
                int(k): v for k, v in metadata['chunk_to_doc_mapping'].items()
            }
            
            # Load TF-IDF
            tfidf_path = f"{index_path}_tfidf.pkl"
            if os.path.exists(tfidf_path):
                import pickle
                with open(tfidf_path, 'rb') as f:
                    tfidf_data = pickle.load(f)
                    self.tfidf_vectorizer = tfidf_data['vectorizer']
                    self.tfidf_matrix = tfidf_data['matrix']
            
            logger.info(f"Index loaded from {index_path}")
            
        except Exception as e:
            logger.error(f"Error loading index: {e}")

class ContactFormFiller:
    """Helper class to extract contact form information from documents."""
    
    def __init__(self):
        self.contact_patterns = {
            'phone': [
                r'(\+?254[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{3})',  # Kenyan format
                r'(\d{4}[-.\s]?\d{3}[-.\s]?\d{3})',  # Short format
                r'(0\d{3}[-.\s]?\d{3}[-.\s]?\d{3})',  # Local format
            ],
            'email': [
                r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
            ],
            'office_hours': [
                r'office hours?[:\s]*([^\n]*(?:\d{1,2}:\d{2}[^\n]*){1,2})',
                r'open[:\s]*([^\n]*(?:\d{1,2}:\d{2}[^\n]*){1,2})',
                r'hours?[:\s]*([^\n]*(?:monday|tuesday|wednesday|thursday|friday)[^\n]*)'
            ],
            'address': [
                r'(?:address|located|office)[:\s]*([^\n]*(?:nairobi|mombasa|kisumu|nakuru)[^\n]*)',
                r'(?:p\.?o\.?\s*box\s*\d+[^\n]*)',
            ]
        }
    
    def extract_contact_info(self, documents: List[Dict]) -> Dict[str, Any]:
        """Extract contact information from documents for form filling."""
        contact_info = {
            'phones': [],
            'emails': [],
            'office_hours': [],
            'addresses': [],
            'departments': []
        }
        
        for doc in documents:
            content = doc.get('content', '') + ' ' + doc.get('snippet', '')
            
            # Extract phone numbers
            for pattern in self.contact_patterns['phone']:
                phones = re.findall(pattern, content, re.IGNORECASE)
                contact_info['phones'].extend(phones)
            
            # Extract emails
            for pattern in self.contact_patterns['email']:
                emails = re.findall(pattern, content, re.IGNORECASE)
                contact_info['emails'].extend(emails)
            
            # Extract office hours
            for pattern in self.contact_patterns['office_hours']:
                hours = re.findall(pattern, content, re.IGNORECASE)
                contact_info['office_hours'].extend(hours)
            
            # Extract addresses
            for pattern in self.contact_patterns['address']:
                addresses = re.findall(pattern, content, re.IGNORECASE)
                contact_info['addresses'].extend(addresses)
            
            # Extract department info
            department_keywords = ['department', 'office', 'ministry', 'authority', 'commission']
            for keyword in department_keywords:
                dept_pattern = rf'{keyword}\s+of\s+([^.\n]*)'
                departments = re.findall(dept_pattern, content, re.IGNORECASE)
                contact_info['departments'].extend(departments)
        
        # Clean and deduplicate
        for key in contact_info:
            contact_info[key] = list(set(contact_info[key]))[:5]  # Limit to 5 items each
        
        return contact_info
    
    def format_for_nairobi_contact_form(self, contact_info: Dict[str, Any], query: str) -> Dict[str, str]:
        """Format extracted info for Nairobi County contact form."""
        
        # Determine subject/category from query
        subject = self._determine_subject(query)
        
        # Pick best contact info
        phone = contact_info['phones'][0] if contact_info['phones'] else "0709 704 704"
        email = contact_info['emails'][0] if contact_info['emails'] else ""
        
        # Generate message based on query and found info
        message_parts = [f"Inquiry regarding: {query}"]
        
        if contact_info['office_hours']:
            message_parts.append(f"Office hours: {contact_info['office_hours'][0]}")
        
        if contact_info['departments']:
            message_parts.append(f"Related department: {contact_info['departments'][0]}")
        
        message = "\n".join(message_parts)
        
        return {
            'subject': subject,
            'name': 'CivicNavigator User',  # Could be filled from user profile
            'email': email or 'user@example.com',  # User's email
            'phone': phone,
            'message': message,
            'department': self._map_to_county_department(query),
            'priority': self._determine_priority(query)
        }
    
    def _determine_subject(self, query: str) -> str:
        """Determine email subject from query."""
        query_lower = query.lower()
        
        if any(term in query_lower for term in ['waste', 'garbage', 'collection']):
            return "Waste Management Inquiry"
        elif any(term in query_lower for term in ['business', 'permit', 'license']):
            return "Business Permit Inquiry"
        elif any(term in query_lower for term in ['water', 'connection', 'sewer']):
            return "Water Services Inquiry"
        elif any(term in query_lower for term in ['road', 'transport', 'matatu']):
            return "Transport/Roads Inquiry"
        elif any(term in query_lower for term in ['health', 'hospital', 'clinic']):
            return "Health Services Inquiry"
        else:
            return "General Inquiry"
    
    def _map_to_county_department(self, query: str) -> str:
        """Map query to appropriate county department."""
        query_lower = query.lower()
        
        department_map = {
            'Environment & Waste Management': ['waste', 'garbage', 'environment', 'recycling'],
            'Trade & Business': ['business', 'permit', 'license', 'trade'],
            'Water & Sanitation': ['water', 'sewer', 'sanitation', 'drainage'],
            'Transport & Roads': ['road', 'transport', 'traffic', 'parking'],
            'Health Services': ['health', 'hospital', 'clinic', 'medical'],
            'Education': ['school', 'education', 'bursary'],
            'Planning & Development': ['building', 'construction', 'development', 'planning']
        }
        
        for department, keywords in department_map.items():
            if any(keyword in query_lower for keyword in keywords):
                return department
        
        return 'General Administration'
    
    def _determine_priority(self, query: str) -> str:
        """Determine priority level from query."""
        query_lower = query.lower()
        
        high_priority_terms = ['emergency', 'urgent', 'broken', 'not working', 'complaint']
        medium_priority_terms = ['application', 'apply', 'process', 'how to']
        
        if any(term in query_lower for term in high_priority_terms):
            return 'High'
        elif any(term in query_lower for term in medium_priority_terms):
            return 'Medium'
        else:
            return 'Normal'

# Usage example and testing
async def test_kenyan_retrieval_system():
    """Test the enhanced Kenyan retrieval system."""
    
    retriever = KenyanDocumentRetriever(
        kb_path="./knowledge_base/sourced_pdfs",
        index_path="./indexes/kenyan_kb"
    )
    
    # Initialize system (this will process local PDFs and fetch remote URLs)
    print("Initializing Kenyan document retrieval system...")
    await retriever.initialize_full_system(rebuild_index=True)
    
    # Test queries specific to Nairobi County
    test_queries = [
        "When is garbage collected in South C ward?",
        "How do I apply for a business permit in Nairobi County?",
        "What are the requirements for water connection?",
        "How to report a broken streetlight in Karen?",
        "Nairobi County service charter customer service hours",
        "Environmental best practices for waste management",
        "County assembly schedule and meetings",
        "Solid waste management act penalties"
    ]
    
    contact_filler = ContactFormFiller()
    
    print("\n" + "="*80)
    print("TESTING KENYAN GOVERNMENT DOCUMENT RETRIEVAL")
    print("="*80)
    
    for query in test_queries:
        print(f"\nQuery: {query}")
        print("-" * 60)
        
        # Search documents
        results = retriever.search(query, top_k=3)
        
        if results:
            for i, result in enumerate(results, 1):
                print(f"{i}. {result['title']}")
                print(f"   Score: {result['score']:.3f}")
                print(f"   Category: {result['category']}")
                print(f"   County: {result['county']}")
                print(f"   Document Type: {result['document_type']}")
                print(f"   Snippet: {result['snippet'][:200]}...")
                if result['source_url']:
                    print(f"   Source: {result['source_url']}")
                print()
            
            # Extract contact info for form filling
            contact_info = contact_filler.extract_contact_info(results)
            form_data = contact_filler.format_for_nairobi_contact_form(contact_info, query)
            
            print("Contact Form Data:")
            for key, value in form_data.items():
                print(f"   {key}: {value}")
        else:
            print("   No results found")
        
        print("\n" + "="*60)

if __name__ == "__main__":
    # Run the test
    asyncio.run(test_kenyan_retrieval_system())